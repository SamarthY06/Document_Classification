import os
import PyPDF2
from PIL import Image
from docx import Document
from typing import List
from extractor import Extractor
from pdf2image import convert_from_path

import pytesseract

#poppler_path=r'C:\Users\Samarth\Downloads\Release-23.01.0-0\poppler-23.01.0\Library\bin\pdfimages.exe'
#poppler_path=r'C:\Program Files\poppler-0.68.0\bin'
class DocumentSplitter:

    def __init__(self):
        pass

    def pdf_splitter(self, file_path: str) -> List[str]:
        """
        Accepts the file path of a pdf and returns a list of pdf files containing
        individual documents
        """
        documents = []
        pages = convert_from_path(file_path,500,poppler_path=r'C:\poppler-0.68.0\bin')
        for i, page in enumerate(pages):
            extractor = Extractor()
            document_text = extractor.pdf_extractor(page)
            page.save("data/test_split/predict/page_{}.jpg".format(i), "JPEG")
            document_name = "page_{}.pdf".format(i)
            documents.append({"path":document_name,"text":document_text})
        return documents

        
    def image_splitter(self, file_path: str) -> List[str]:
        """
        Accepts the file path of an image and returns a list of image files containing
        individual documents
        """
        # Open the image file
        with Image.open(file_path) as img:
            width, height = img.size
            # Split the image into individual documents
            document_count = height // (width // 2)  # Assumes 2 documents per row
            documents = []
            for i in range(document_count):
                left = 0
                upper = i * (width // 2)
                right = width
                lower = (i + 1) * (width // 2)
                document = img.crop((left, upper, right, lower))
                # Save the document to a new image file
                document_name = "img_document_{}.jpg".format(i)
                document.save(f'data/test_split/predict/{document_name}')
                # Get the text from the image using OCR
                extractor = Extractor()
                document_text = extractor.img_extractor(document)
                documents.append({"path":document_name,"text":document_text})
            return documents

    def docx_splitter(self, file_path: str) -> List[str]:
        """
        Accepts the file path of a docx file and returns a list of """
        # Open the docx file
        doc = Document(file_path)
        # Split the docx into individual documents
        document_count = len(doc.paragraphs) // 10 # Assumes 10 paragraphs per document
        documents = []
        for i in range(document_count):
            document = Document()
            document_text = ""
        for j in range(i * 10, (i + 1) * 10):
            document.add_paragraph(doc.paragraphs[j].text)
            document_text += doc.paragraphs[j].text
        # Save the document to a new docx file
        document_path = "word_document_{}.docx".format(i)
        document_img =  "word_document_{}.jpg".format(i)
        document.save(f'data/test_split/predict/{document_img}')
        documents.append({"path":document_path,"text":document_text})
        return documents

# pdf_path = 'data/test/predict/test_1.pdf'
# c = DocumentSplitter()
# documenents = c.pdf_splitter(pdf_path)
# print(documenents)


